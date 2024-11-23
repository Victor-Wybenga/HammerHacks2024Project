import customtkinter as ctk
import pathlib
import ParseQuiz, Quiz
import docx

def upload_document() -> Quiz.Quiz:
    filename = ctk.filedialog.askopenfilename(filetypes=[("Word Document", "*.docx")])
    quiz_document = ParseQuiz.QuizDocument(pathlib.Path(filename))
    parsed_quiz = quiz_document.make_quiz()
    if isinstance(parsed_quiz, ParseQuiz.FormatError):
        print("Format Error!")
        quit()
    return Quiz.Quiz(parsed_quiz, 30)

def main() -> None:
    doc = docx.Document()
    quiz = upload_document()
    doc.add_paragraph(quiz.topic, style='Title')
    for idx, question in enumerate(quiz.questions):
        doc.add_paragraph(f"{idx + 1}: {question.question_text}")
        if isinstance(question, Quiz.MultipleChoiceQuestion):
            doc.add_paragraph("\n".join([f"{chr(idx + 65)} : {choice}" for idx, choice in enumerate(question.choices)]))
        if isinstance(question, Quiz.TrueFalseQuestion):
            doc.add_paragraph("A. True\nB. False")
        if idx % 6 == 5:
            doc.add_page_break()
    
    doc.add_page_break()
    doc.add_paragraph("Answers", style='Title')
    for idx, question in enumerate(quiz.questions):
        doc.add_paragraph(f"{idx + 1}: {question.answer}")
    f = ctk.filedialog.asksaveasfile(mode='wb', defaultextension=".docx")
    if f == None:
        return
    doc.save(f)

if __name__ == "__main__":
    main()